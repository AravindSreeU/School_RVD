
import random as rd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH as align
from docx.shared import Pt, Inches
import tkinter as tk
from tkinter import ttk

# Required functions for generating worksheet
def add_qs(worksheet, m, n, no_of_problems):  # Function for generating Addition problems
    for i in range(1,no_of_problems+1):
        num1 = rd.randrange(m, n)
        num2 = rd.randrange(m, n)
        text = f"{i}.   {num1}\n   + {num2}\n    ________\n    ________\n"
        paragraph = worksheet.add_paragraph(text)

def sub_qs(worksheet, m, n, no_of_problems):  # Function for generating Subtraction problems
    for i in range(1, no_of_problems + 1):
        num1 = rd.randrange(m, n)
        num2 = rd.randrange(m, num1 + 1)  # Ensure num2 is less than or equal to num1
        text = f"{i}.   {num1}\n   - {num2}\n    ________\n    ________\n"
        paragraph = worksheet.add_paragraph(text)

def mul_qs(worksheet, m, n, no_of_problems):  # Functions for generating Multiplication problems
    for i in range(1,no_of_problems+1):
        num1 = rd.randrange(m, n)
        num2 = rd.randrange(m, n)
        text = f"{i}.   {num1}\n   x {num2}\n    ________\n    ________\n"
        paragraph = worksheet.add_paragraph(text)

#Function to declare the operation in the final step
def generate_problems(operation, m, n, no_of_problems):
    if operation == "Addition":
        add_qs(worksheet, m, n, no_of_problems)
    elif operation == "Subtraction":
        sub_qs(worksheet, m, n, no_of_problems)
    elif operation == "Multiplication":
        mul_qs(worksheet, m, n, no_of_problems)

#Function to define the operation to be performed after the button is clicked
def on_select():
    global operation_selected
    operation_selected = var.get()

#Function to update the values of m and n inside the problems generating part
def update_m_n():
    global m, n
    no_of_digits = int(entry_no_of_digits.get())
    if no_of_digits == 1:
        m, n = 1, 10
    elif no_of_digits == 2:
        m, n = 10, 100
    elif no_of_digits == 3: 
        m, n = 100, 1000
    else:
        m, n = 1000, 10000

# Creating a document
worksheet = Document()

# Document title and alignment
title = worksheet.add_heading("Customizable Title", 0)
title.alignment = align.CENTER

# Change the font style of the title
title_font = title.runs[0].font
title_font.name = 'Arial'  # Change to your desired font
title_font.size = Pt(24)  # Change to your desired font size

# Building main window
root = tk.Tk()
root.geometry("400x350")
root.resizable(False, False)
root.title("Mathwork")

# Labels and input section
file_name = tk.Label(root, text="1. Enter the file name:")  # file name
file_name.grid(row=0, column=0, padx=5, pady=5, sticky='w')
entry_file_name = tk.Entry(root)
entry_file_name.grid(row=0, column=1, padx=10, pady=5)

# Select the arithmetic operation by radio button
operation = tk.Label(root, text="2. Choose an operation:")  # operation
operation.grid(row=1, column=0, padx=10, pady=10, sticky='w')
var = tk.StringVar()
operation_selected = None
radio1 = ttk.Radiobutton(root, text="Addition", variable=var, value="Addition", command=on_select)
radio2 = ttk.Radiobutton(root, text="Subtraction", variable=var, value="Subtraction", command=on_select)
radio3 = ttk.Radiobutton(root, text="Multiplication", variable=var, value="Multiplication", command=on_select)
radio1.grid(row=2, column=0, padx=10, pady=5, sticky="w")
radio2.grid(row=3, column=0, padx=10, pady=5, sticky="w")
radio3.grid(row=4, column=0, padx=10, pady=5, sticky="w")

# To input number of digits of the problems
no_of_digits = tk.Label(root, text='3. Enter number of digits:')
no_of_digits.grid(row=5, column=0, padx=10, pady=5, sticky='w')
entry_no_of_digits = tk.Entry(root)
entry_no_of_digits.grid(row=5, column=1, padx=10, pady=5)

# To input the number of problems needed
num_problems = tk.Label(root, text="4. Enter the number of problems:")
num_problems.grid(row=6, column=0, padx=10, pady=5, sticky='w')
entry_num_problems = tk.Entry(root)
entry_num_problems.grid(row=6, column=1, padx=10, pady=5)

file_path = tk.Label(root, text="5. Enter the file path:")
file_path.grid(row=7, column=0, padx=10, pady=5, sticky='w')
entry_file_path = tk.Entry(root)
entry_file_path.grid(row=7, column=1, padx=10, pady=5, sticky='e')

# Button to generate worksheet
btn_generate = tk.Button(root, text='Generate', command=lambda: [update_m_n(), generate_problems(operation_selected, m, n, int(entry_num_problems.get())), worksheet.save(entry_file_path.get() + '/' + entry_file_name.get()+'.docx')])
btn_generate.grid(row=8, column=0, padx=10, pady=5, sticky='w')
root.mainloop()
